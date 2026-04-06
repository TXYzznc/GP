#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从 .docx 文件中提取格式信息
"""

from docx import Document
from docx.shared import Pt, RGBColor
import json

def extract_format(docx_path):
    """提取 .docx 文件的格式信息"""
    doc = Document(docx_path)
    
    format_info = {
        "page": {},
        "font": {},
        "paragraph": {},
        "styles": {}
    }
    
    # 提取页面设置
    section = doc.sections[0]
    format_info["page"] = {
        "top_margin": section.top_margin.cm,
        "bottom_margin": section.bottom_margin.cm,
        "left_margin": section.left_margin.cm,
        "right_margin": section.right_margin.cm,
    }
    
    # 提取样式信息
    for style in doc.styles:
        if style.type == 1:  # 段落样式
            style_info = {
                "name": style.name,
                "font": {},
                "paragraph": {}
            }
            
            if style.font.name:
                style_info["font"]["name"] = style.font.name
            if style.font.size:
                style_info["font"]["size"] = style.font.size.pt
            if style.font.bold is not None:
                style_info["font"]["bold"] = style.font.bold
            
            if style.paragraph_format.line_spacing:
                style_info["paragraph"]["line_spacing"] = style.paragraph_format.line_spacing
            
            format_info["styles"][style.name] = style_info
    
    # 提取正文格式
    if doc.paragraphs:
        first_para = doc.paragraphs[0]
        if first_para.runs:
            run = first_para.runs[0]
            if run.font.name:
                format_info["font"]["default_name"] = run.font.name
            if run.font.size:
                format_info["font"]["default_size"] = run.font.size.pt
    
    return format_info

if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("用法: python extract_docx_format.py <docx_file>")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    
    try:
        format_info = extract_format(docx_file)
        print(json.dumps(format_info, indent=2, ensure_ascii=False))
    except Exception as e:
        print(f"错误: {e}", file=sys.stderr)
        sys.exit(1)

#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
论文 .docx 生成脚本
使用 python-docx 库将 Markdown 章节转换为 Word 文档

使用方法: python generate_thesis_docx.py

依赖: pip install python-docx
"""

import os
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ 缺少依赖: python-docx")
    print("请运行: pip install python-docx")
    sys.exit(1)

# 配置
CONFIG = {
    'title': 'Clash of Gods 游戏系统设计与实现',
    'author': '学生姓名',
    'date': datetime.now().strftime('%Y-%m-%d'),
    'font_body': '宋体',
    'font_heading': '黑体',
    'font_size_body': 12,
    'font_size_h1': 18,
    'font_size_h2': 14,
    'font_size_h3': 12,
    'line_spacing': 1.5,
}

# 章节文件列表
CHAPTERS = [
    '第1章_绪论.md',
    '第2章_相关工作与技术基础.md',
    '第3章_系统总体设计与架构.md',
    '第4章_战斗系统.md',
    '第5章_物品系统.md',
    '第6章_卡牌系统.md',
    '第7章_TA系统.md',
    '第8章_性能优化.md',
    '第9章_总结展望.md',
]

def set_cell_border(cell, **kwargs):
    """设置表格单元格边框"""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '12')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), 'CCCCCC')
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)

def parse_markdown(content):
    """解析 Markdown 内容为段落数据"""
    paragraphs = []
    lines = content.split('\n')
    in_code_block = False
    code_block_content = ''
    code_block_language = ''
    
    for line in lines:
        # 代码块处理
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block_language = line[3:].strip()
                code_block_content = ''
            else:
                in_code_block = False
                if code_block_content.strip():
                    paragraphs.append({
                        'type': 'code',
                        'content': code_block_content.strip(),
                        'language': code_block_language,
                    })
            continue
        
        if in_code_block:
            code_block_content += line + '\n'
            continue
        
        # 标题处理
        if line.startswith('# '):
            paragraphs.append({
                'type': 'heading1',
                'content': line[2:].strip(),
            })
            continue
        
        if line.startswith('## '):
            paragraphs.append({
                'type': 'heading2',
                'content': line[3:].strip(),
            })
            continue
        
        if line.startswith('### '):
            paragraphs.append({
                'type': 'heading3',
                'content': line[4:].strip(),
            })
            continue
        
        # 空行处理
        if line.strip() == '':
            if paragraphs and paragraphs[-1]['type'] != 'empty':
                paragraphs.append({'type': 'empty'})
            continue
        
        # 普通段落
        if line.strip():
            paragraphs.append({
                'type': 'normal',
                'content': line.strip(),
            })
    
    return paragraphs

def read_chapter(filename):
    """读取章节文件"""
    filepath = Path(__file__).parent / '章节文件' / filename
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"❌ 无法读取文件: {filepath}")
        return ''

def add_paragraph_to_doc(doc, para_data):
    """添加段落到文档"""
    if para_data['type'] == 'heading1':
        p = doc.add_heading(para_data['content'], level=1)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(CONFIG['line_spacing'] * CONFIG['font_size_body'])
        for run in p.runs:
            run.font.name = CONFIG['font_heading']
            run.font.size = Pt(CONFIG['font_size_h1'])
    
    elif para_data['type'] == 'heading2':
        p = doc.add_heading(para_data['content'], level=2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(CONFIG['line_spacing'] * CONFIG['font_size_body'])
        for run in p.runs:
            run.font.name = CONFIG['font_heading']
            run.font.size = Pt(CONFIG['font_size_h2'])
    
    elif para_data['type'] == 'heading3':
        p = doc.add_heading(para_data['content'], level=3)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(CONFIG['line_spacing'] * CONFIG['font_size_body'])
        for run in p.runs:
            run.font.name = CONFIG['font_heading']
            run.font.size = Pt(CONFIG['font_size_h3'])
    
    elif para_data['type'] == 'code':
        p = doc.add_paragraph(para_data['content'])
        p.style = 'Normal'
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(CONFIG['line_spacing'] * CONFIG['font_size_body'])
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        # 设置背景色
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'F5F5F5')
        p._element.get_or_add_pPr().append(shading_elm)
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
    
    elif para_data['type'] == 'normal':
        p = doc.add_paragraph(para_data['content'])
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(CONFIG['line_spacing'] * CONFIG['font_size_body'])
        for run in p.runs:
            run.font.name = CONFIG['font_body']
            run.font.size = Pt(CONFIG['font_size_body'])
    
    elif para_data['type'] == 'empty':
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(CONFIG['line_spacing'] * CONFIG['font_size_body'])

def generate_thesis():
    """生成最终文档"""
    print('📝 开始生成论文 .docx 文件...\n')
    
    # 创建文档
    doc = Document()
    
    # 设置页面边距 (2.54cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # 添加标题页
    print('📄 添加标题页...')
    title = doc.add_heading(CONFIG['title'], level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = CONFIG['font_heading']
        run.font.size = Pt(CONFIG['font_size_h1'])
    
    author = doc.add_paragraph(f"作者: {CONFIG['author']}")
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in author.runs:
        run.font.name = CONFIG['font_body']
        run.font.size = Pt(CONFIG['font_size_body'])
    
    date_p = doc.add_paragraph(f"日期: {CONFIG['date']}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in date_p.runs:
        run.font.name = CONFIG['font_body']
        run.font.size = Pt(CONFIG['font_size_body'])
    
    doc.add_page_break()
    
    # 添加各章节
    for chapter in CHAPTERS:
        print(f'📖 处理章节: {chapter}')
        content = read_chapter(chapter)
        
        if content:
            paragraphs = parse_markdown(content)
            for para_data in paragraphs:
                add_paragraph_to_doc(doc, para_data)
            doc.add_page_break()
    
    # 保存文档
    output_path = Path(__file__).parent / 'Clash_Of_Gods_毕业论文_最终版.docx'
    print(f'\n💾 保存文档到: {output_path}')
    
    try:
        doc.save(str(output_path))
        file_size = output_path.stat().st_size / 1024
        print(f'\n✅ 论文生成成功！')
        print(f'📊 文档信息:')
        print(f'   - 标题: {CONFIG["title"]}')
        print(f'   - 作者: {CONFIG["author"]}')
        print(f'   - 章节数: {len(CHAPTERS)}')
        print(f'   - 文件大小: {file_size:.2f} KB')
        print(f'   - 输出路径: {output_path}')
    except Exception as e:
        print(f'\n❌ 生成失败: {e}')
        sys.exit(1)

if __name__ == '__main__':
    generate_thesis()

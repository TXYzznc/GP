#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""提取docx文件内容"""

from docx import Document
import sys

def extract_docx(input_path, output_path):
    doc = Document(input_path)
    
    content = []
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text)
    
    # 提取表格
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            content.append(' | '.join(row_text))
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(content))
    
    print(f'已提取 {len(content)} 段落/行')

if __name__ == '__main__':
    if len(sys.argv) >= 3:
        extract_docx(sys.argv[1], sys.argv[2])